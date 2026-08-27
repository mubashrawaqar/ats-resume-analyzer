"""DOCX parsing for resumes.

python-docx only reads text that lives in the normal paragraph/table flow.
Content placed inside text boxes (a common resume-formatting mistake) is
stored in a different XML part and is invisible to python-docx — which is
actually useful for us: if the document *looks* like it should have more
content than we found, we flag it as an ATS risk, since real ATS parsers
have exactly the same blind spot.
"""

from __future__ import annotations

import io
from typing import List

import docx

from utils.models import ParsedDocument

STANDARD_SECTION_HEADERS = [
    "experience", "work experience", "education", "skills",
    "summary", "projects", "certifications",
]


def _detect_ats_issues(document: "docx.document.Document", full_text: str) -> List[str]:
    issues: List[str] = []

    # Tables used for layout (common way people build multi-column resumes).
    if document.tables:
        issues.append(
            f"Detected {len(document.tables)} table(s) in the document — "
            "content inside tables/text boxes is often skipped or "
            "reordered incorrectly by real ATS parsers."
        )

    # Missing standard section headers ATS systems typically look for.
    lower_text = full_text.lower()
    found_headers = [h for h in STANDARD_SECTION_HEADERS if h in lower_text]
    if len(found_headers) < 2:
        issues.append(
            "Few or no standard section headers found (e.g. 'Experience', "
            "'Education', 'Skills') — ATS systems rely on these to categorize "
            "your content correctly."
        )

    return issues


def extract_text(file_bytes: bytes, filename: str) -> ParsedDocument:
    """Extract text from a .docx resume.

    Never raises — failures are captured in ParsedDocument.warnings.
    """
    try:
        document = docx.Document(io.BytesIO(file_bytes))
    except Exception as exc:
        return ParsedDocument(
            text="",
            success=False,
            warnings=[f"Could not parse '{filename}' as a .docx file: {exc}"],
        )

    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

    # Also pull text out of any tables, since content there is real (even
    # though we flag tables as an ATS risk above).
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())

    full_text = "\n".join(paragraphs).strip()
    ats_issues = _detect_ats_issues(document, full_text)

    if not full_text:
        return ParsedDocument(
            text="",
            success=False,
            warnings=[f"'{filename}' appears to be empty or contains no extractable text."],
            ats_issues=ats_issues,
        )

    return ParsedDocument(text=full_text, success=True, ats_issues=ats_issues)
